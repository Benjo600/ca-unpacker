"""Builds CA-Unpacker-Product-Roadmap.docx.

Shares the visual language of create_manual_test_guide.py so the two
documents read as one set. Content is written in plain English for a
reader who is not a programmer.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\Admin\OneDrive\Desktop\CA idea")
OUTPUT = ROOT / "CA-Unpacker-Product-Roadmap.docx"

NAVY = "163A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDECEC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "5C6670"
GRID = "CDD5DF"
WHITE = "FFFFFF"
BLACK = "1F2328"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        if edge in kwargs:
            tag = "w:" + edge
            element = tc_mar.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_mar.append(element)
            element.set(qn("w:w"), str(kwargs[edge]))
            element.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=BLACK, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_callout(doc, label, text, fill=PALE_GOLD, color=BLACK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=fill, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_body(doc, text, size=11, after=6, italic=None, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, italic=italic, color=color)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=10.7, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.7)
    return p


def add_numbered(doc, index, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.32)
    r = p.add_run(f"{index}.  ")
    set_run_font(r, size=10.7, bold=True, color=BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.7)
    return p


def add_check(doc, text):
    p = doc.add_paragraph(style="Checklist")
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run("[  ]   " + text)
    set_run_font(r, size=10.7)
    return p


def add_mini_heading(doc, text, color=NAVY, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run_font(r, size=9.5, color=color, bold=True)
    return p


def add_table(doc, headers, rows, widths, header_fill=NAVY, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    head = table.rows[0]
    set_repeat_table_header(head)
    for cell, label in zip(head.cells, headers):
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10, color=WHITE, bold=True)
    for idx, values in enumerate(rows):
        row = table.add_row()
        for cidx, (cell, value) in enumerate(zip(row.cells, values)):
            if zebra and idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=10, bold=(cidx == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_stage_banner(doc, number, title, status, status_color):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [7100, 2260])
    set_table_borders(table, color=NAVY, size="8")
    left = table.cell(0, 0)
    set_cell_shading(left, NAVY)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(number.upper() + "\n")
    set_run_font(r, size=9, color="A9C2DA", bold=True)
    r = p.add_run(title)
    set_run_font(r, size=15, color=WHITE, bold=True)
    right = table.cell(0, 1)
    set_cell_shading(right, NAVY)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(status.upper())
    set_run_font(r, size=10, color=status_color, bold=True)
    set_table_geometry(table, [7100, 2260])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


DONE_STAGES = [
    (
        "Stage 1",
        "Windows shell",
        "Prove this is a real Windows program, not a website.",
        "A CA opens an app from the Start menu, adds a client, closes it, opens it again, and the client is "
        "still there. No internet, no Docker, no browser tab.",
        "Solid. Nothing to revisit.",
    ),
    (
        "Stage 2",
        "Dump and router",
        "Prove a CA can dump a messy folder without labelling anything.",
        "Drag a folder of mixed files onto a client's month. Each file gets recognised as a bank statement, "
        "an invoice, a GST file, a Tally export, and so on. Anything it cannot place goes to a Needs Review "
        "list where the CA can correct it by hand.",
        "Works. Two gaps: older .xls files are refused outright, and there is still no way to delete a file, "
        "a client, or a month once it has been added.",
    ),
    (
        "Stage 3",
        "Bank pack",
        "Prove the core demo - a messy bank PDF becomes a clean, checked Excel, entirely offline.",
        "This is the first stage a CA would actually pay attention to. Transactions come out with date, "
        "description, debit, credit and balance, and the app adds up the running balance to check its own work.",
        "This is the one stage proven on real statements, for HDFC, ICICI, SBI, Axis and Kotak. It has two "
        "honesty problems, both fixed in Stage 3.5 next: rows that do not add up can be dropped instead of "
        "flagged, and the balance check can pass without really checking anything.",
    ),
]

# Stages that exist in code but do not work well enough on real documents yet.
NEEDS_WORK_STAGES = [
    (
        "Stage 4",
        "Trust - crops, scans and passwords",
        "A CA can verify any number without trusting the software.",
        "Click an amount in the app and see a cropped picture of that exact spot on the original page. "
        "Scanned statements are read by text-recognition software bundled inside the app, so nothing is "
        "uploaded. Password-protected PDFs ask for the password once.",
        [
            "Reading scans is the weak point. Text recognition on a real photographed or faxed statement is "
            "far less reliable than on a clean digital PDF, and the results have not been measured against "
            "real samples.",
            "The password is forgotten the moment the app closes, so a protected statement asks again on every "
            "restart and after every re-run.",
            "The part of the app that produces the crops is called through a hundred lines of guesswork that "
            "try several different ways of asking until one works. When it fails, the real reason is hidden.",
            "The bundled text-recognition software is not installed in a predictable place, so whether scans "
            "work at all depends on the machine.",
        ],
        [
            "Measure scan accuracy on at least ten real scanned statements and write the number down.",
            "Keep the password safely for as long as the app is open.",
            "Replace the guesswork with a direct call, so failures report their real cause.",
            "Make the text-recognition software install to one known location, every time.",
        ],
        "About one to two weeks.",
    ),
    (
        "Stage 5",
        "Invoices",
        "A folder of purchase bills becomes a GST-aware register.",
        "Printed bills, as PDFs or photos, become rows with supplier, GSTIN, invoice number, date, tax and "
        "total. Bad GSTINs and odd HSN codes get flagged rather than silently corrected.",
        [
            "This is the hardest reading problem in the whole product and the least proven. Every supplier "
            "lays their bill out differently, and the app has only been tested against invoices it generated "
            "itself - which naturally all look the same.",
            "A bill that produces no line items is thrown out entirely and marked as failed, even when the "
            "totals were read correctly. A partly-read bill is more useful than no bill.",
            "There is no measure of how often it gets a bill right, so there is no way to tell whether a "
            "change made things better or worse.",
        ],
        [
            "Collect fifty real purchase bills from different suppliers and measure how many are read correctly.",
            "Keep a bill whose totals were read even when the line items could not be, and flag it clearly.",
            "Fix the most common layout failures that measurement exposes, then measure again.",
            "Set a number the stage must beat before it counts as done.",
        ],
        "About two to three weeks. Most of this is measurement and patient fixing, not new invention.",
    ),
    (
        "Stage 6",
        "GST portal files",
        "The ugly files downloaded from the GST portal become sheets a CA will actually open.",
        "The official GSTR-1, GSTR-2B and GSTR-3B downloads are unreadable raw data files. The app turns them "
        "into formatted Excel sheets in one step, with no portal login and no internet.",
        [
            "The reader was built against invented sample files. Real portal downloads carry sections, "
            "corrections and note types that those samples never contained.",
            "This is the most fixable of the four, because the portal's format is published and consistent - "
            "it simply has more to it than has been handled so far.",
            "Getting this right matters more than it looks: Stage 8 compares everything against GSTR-2B, so a "
            "section missed here becomes a wrong answer there.",
        ],
        [
            "Run real GSTR-1, GSTR-2B and GSTR-3B downloads from at least three different clients through it.",
            "Handle the sections currently missing, including credit and debit notes and amendments.",
            "Check the totals produced against the summary the portal itself prints.",
        ],
        "About one week, and the most certain of the four.",
    ),
    (
        "Stage 7",
        "Books exports",
        "Accounting books enter the pipeline without opening Tally or Zoho.",
        "A Tally or Zoho export becomes purchase and sales registers in the same shape as everything else - "
        "which is precisely what makes Stage 8 possible.",
        [
            "Recognising these files is far too trusting. Any spreadsheet with a column called 'invoice "
            "number' is treated as a Zoho export, so unrelated files get read as books and produce nonsense "
            "rows.",
            "Older .xls files are refused outright, and that is the format Tally and several bank portals "
            "still produce by default.",
            "Tally exports vary by version and by what the user chose to export; only one shape has been "
            "tested.",
            "Deciding whether an entry is a purchase or a sale falls back to guessing when the export does "
            "not say, and a wrong guess puts an entry in the wrong register.",
        ],
        [
            "Recognise these files on more than a single column name.",
            "Accept older .xls files.",
            "Test against exports from at least two Tally versions and a real Zoho account.",
            "When purchase or sale genuinely cannot be determined, flag it instead of guessing.",
        ],
        "About one to two weeks.",
    ),
]


def stage_section(doc, banner, title, status, status_color, one_line, why, builds,
                  done_checks, timing, not_this, closing=None, closing_fill=PALE_BLUE,
                  closing_label=None):
    add_stage_banner(doc, banner, title, status, status_color)
    add_mini_heading(doc, "In one line", before=4)
    add_body(doc, one_line, size=11, after=6)
    add_mini_heading(doc, "Why it matters")
    for para in why:
        add_body(doc, para, size=10.7, after=5)
    add_mini_heading(doc, "What gets built")
    for idx, item in enumerate(builds, start=1):
        add_numbered(doc, idx, item)
    add_mini_heading(doc, "How we know it is done")
    for item in done_checks:
        add_check(doc, item)
    add_mini_heading(doc, "Rough size")
    add_body(doc, timing, size=10.7, after=5)
    add_mini_heading(doc, "Not in this stage")
    add_body(doc, not_this, size=10.7, after=6)
    if closing:
        add_callout(doc, closing_label or "Note:", closing, fill=closing_fill)


def build_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCT ROADMAP")
    set_run_font(r, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CA Unpacker")
    set_run_font(r, size=34, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("From working prototype to a product a firm can pay for")
    set_run_font(r, size=14, color=DARK_BLUE)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2400, 6960])
    set_table_borders(meta, color="D8E0E8")
    for row, (label, value) in zip(
        meta.rows,
        [
            ("Prepared", "27 August 2026"),
            ("Where we are", "Stages 1 to 3 hold up. Stages 4 to 7 are built but do not work well enough yet."),
            ("What is left", "Repair Stage 3, make Stages 4 to 7 reliable, then build Stages 8, 9 and 10."),
            ("Rough time to 1.0", "About 4 to 6 months of focused solo work."),
        ],
    ):
        set_cell_shading(row.cells[0], PALE_BLUE)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10)
    set_table_geometry(meta, [2400, 6960])

    doc.add_paragraph()
    add_callout(
        doc,
        "Read this first:",
        "The foundations are solid and the bank pack genuinely works. But four stages that were recorded as "
        "passed - invoices, GST files, books exports and the trust features - do not yet work well enough to "
        "put in front of a CA. Making those four reliable is the real job, and it comes before any new feature.",
        fill=PALE_RED,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("Written in plain English. No programming knowledge needed.")
    set_run_font(r, size=11, color=MID_GRAY, italic=True)
    doc.add_page_break()


def build_how_to_read(doc):
    doc.add_heading("How to read this roadmap", level=1)
    add_body(
        doc,
        "The app is built in stages. A stage is one complete promise to the user - not a list of chores. "
        "You do not start the next stage until the current one truly works.",
    )
    add_body(doc, "Every stage below is written the same way, so you can compare them at a glance:")
    add_bullet(doc, "the one thing this stage promises, in a single sentence.", bold_lead="In one line - ")
    add_bullet(
        doc,
        "why a chartered accountant would care. If there is no answer here, the stage is not worth building.",
        bold_lead="Why it matters - ",
    )
    add_bullet(doc, "the actual work, in order.", bold_lead="What gets built - ")
    add_bullet(
        doc,
        "the test that proves it. If every box ticks, the stage is done. If one fails, it is not.",
        bold_lead="How we know it is done - ",
    )
    add_bullet(
        doc,
        "work that belongs to a later stage. Doing it early is a mistake, not a bonus.",
        bold_lead="Not in this stage - ",
    )

    add_mini_heading(doc, "The one rule")
    add_callout(
        doc,
        "The rule:",
        "Do not start the next stage until the current stage's checklist is fully ticked. A half-working stage "
        "is a failed stage. If a later stage needs something fixed further back, go back and fix it first.",
        fill=PALE_BLUE,
    )

    add_mini_heading(doc, "About the time estimates")
    add_body(
        doc,
        "The times below are rough, and assume one person working on this steadily - not full-time employment, "
        "but real focused hours. Treat them as a sense of proportion, not a promise. Stage 8 is genuinely bigger "
        "than Stage 9, and that comparison is the useful part.",
        size=10.7,
    )
    doc.add_page_break()


def build_status(doc):
    doc.add_heading("Where the project stands today", level=1)
    add_body(
        doc,
        "Three of the ten planned stages hold up. A CA can install the app, drop a messy folder of a client's "
        "paperwork onto it, and get a clean, checked bank statement back as Excel - without the files ever "
        "leaving their computer. That part is real, and the hard plumbing underneath it is done.",
    )
    add_body(
        doc,
        "Four more stages - the trust features, invoices, GST portal files and books exports - have been built "
        "and have code behind them, but they do not work well enough yet. They pass the app's own test files "
        "and fall down on real ones. Treating them as finished is what makes the current picture look better "
        "than it is.",
    )
    add_body(
        doc,
        "So the work ahead has three parts: repair one honesty problem in the bank reader, make those four "
        "stages genuinely reliable, and only then build the three stages that were never started.",
    )

    add_mini_heading(doc, "Stage status at a glance")
    add_table(
        doc,
        ["Stage", "What it gives a CA", "Status"],
        [
            ("1  Windows shell", "A real desktop app that remembers clients", "Done"),
            ("2  Dump and router", "Drop a messy folder; files get sorted by type", "Done"),
            ("3  Bank pack", "A messy bank PDF becomes a checked Excel", "Done"),
            ("3.5  Trust repair", "The numbers can never be quietly wrong", "NEXT - new"),
            ("4  Trust", "Click any number and see it on the original page", "Needs work"),
            ("5  Invoices", "Purchase bills become a flagged register", "Needs work"),
            ("6  GST portal files", "Ugly 2B / 1 / 3B files become readable sheets", "Needs work"),
            ("7  Books exports", "Tally and Zoho exports become registers", "Needs work"),
            ("8  Master grid", "Books vs 2B vs bank, side by side", "Not started"),
            ("9  License", "Customers can pay, files still never upload", "Not started"),
            ("10  Firm-ready", "A real firm can live on it all year", "Not started"),
        ],
        [2100, 5160, 2100],
    )

    add_callout(
        doc,
        "A note on the records:",
        "The gate log at the back of STAGES.md still records Stages 4 to 7 as passed on 15 August. That is "
        "optimistic - they were signed off against the app's own invented test files, and none of the 49 gate "
        "checkboxes in that document were ever actually ticked. This roadmap reflects how the stages behave on "
        "real documents, which is the only measure that matters. Where the two disagree, trust this one.",
        fill=PALE_GOLD,
    )

    add_callout(
        doc,
        "Why there is a Stage 3.5:",
        "It was not in the original plan. A close read of the code found that the bank reader can silently drop "
        "a transaction it cannot make sense of, and then show a green 'balance matches' tick on what is left. "
        "Everything later compares numbers against each other - so if the numbers going in are quietly "
        "incomplete, the comparisons produce confident nonsense. The repair has to come first.",
        fill=PALE_RED,
    )
    doc.add_page_break()


def build_done_part(doc):
    doc.add_heading("Part 1 - The three stages that hold up", level=1)
    add_body(
        doc,
        "These work, on real files, today. They are the foundation everything else sits on, and none of them "
        "needs revisiting - with one exception noted under Stage 3.",
    )
    for number, title, goal, detail, state in DONE_STAGES:
        add_stage_banner(doc, number + "  -  Done", title, "Complete", "9FD3A8")
        add_mini_heading(doc, "The promise", before=4)
        add_body(doc, goal, size=10.7, after=4)
        add_mini_heading(doc, "In plain words")
        add_body(doc, detail, size=10.7, after=4)
        add_mini_heading(doc, "Where it stands now")
        add_body(doc, state, size=10.7, after=10)
    doc.add_page_break()


def build_needs_work_part(doc):
    doc.add_heading("Part 3 - The four stages that need to be made reliable", level=1)
    add_body(
        doc,
        "These four have code behind them and were recorded as passed. In practice they do not work well "
        "enough to put in front of a CA. They are not starting from nothing - which is good news, because "
        "repairing something that half works is far quicker than building it twice.",
    )
    add_callout(
        doc,
        "The pattern behind all four:",
        "Each was signed off against test files the app invented for itself. Invented files are tidy, "
        "consistent, and quietly built to match what the reader already expects. Real client documents are "
        "none of those things. Nothing below is a surprise - it is what always happens when software is tested "
        "only against its own assumptions.",
        fill=PALE_GOLD,
    )
    add_body(
        doc,
        "The weaknesses listed under each stage were found by reading the code, not by watching the app fail "
        "on your files. You will know things this list does not - correct it where it is wrong.",
        size=10.7,
        italic=True,
        color=MID_GRAY,
    )

    for number, title, promise, detail, problems, fixes, timing in NEEDS_WORK_STAGES:
        add_stage_banner(doc, number + "  -  Needs work", title, "Built, unreliable", "F5D89A")
        add_mini_heading(doc, "What it is supposed to do", before=4)
        add_body(doc, promise, size=10.7, after=4)
        add_body(doc, detail, size=10.7, after=4)
        add_mini_heading(doc, "Why it does not work well enough")
        for item in problems:
            add_bullet(doc, item)
        add_mini_heading(doc, "What would make it reliable")
        for idx, item in enumerate(fixes, start=1):
            add_numbered(doc, idx, item)
        add_mini_heading(doc, "Rough size")
        add_body(doc, timing, size=10.7, after=10)

    add_callout(
        doc,
        "One rule for all four:",
        "Do not mark any of these done again on invented files. Each one needs a number measured against real "
        "documents - how many bills read correctly, how many scans, how many exports - and that number written "
        "into the gate log. A date with no number behind it is how this situation came about.",
        fill=PALE_RED,
    )
    doc.add_page_break()


def build_stage_35(doc):
    stage_section(
        doc,
        "Stage 3.5  -  Do this first",
        "Make the numbers trustworthy",
        "Next up",
        "F5B7B7",
        "Make it impossible for the app to lose a transaction quietly, or to show a green tick it has not earned.",
        [
            "Right now, when the bank reader meets a line whose balance does not fit the running total - a "
            "misread digit, a smudged scan - it throws that line away. The transaction simply is not in the "
            "Excel. Nothing tells the CA it happened.",
            "It gets worse when combined with the balance check. Once the awkward rows are gone, the remaining "
            "rows add up perfectly, so the app reports 'balance matches' in green. Separately, if the app cannot "
            "find the opening and closing balance printed on the statement, it works them out from the rows it "
            "already has - and then checks those rows against themselves. That check always passes.",
            "So the worst case is not a crash. It is a CA looking at a confident green tick on a register that "
            "is missing three transactions. For accounting software, that is the most serious kind of failure "
            "there is.",
        ],
        [
            "Stop throwing rows away. A transaction that does not fit the running balance still goes into the "
            "Excel, marked clearly as needing a human eye.",
            "Count everything. If the app sees forty lines that look like transactions and only produces "
            "thirty-eight rows, it must say so out loud - on screen and in the Needs Review sheet.",
            "Stop giving unearned green ticks. If the opening or closing balance could not be found on the "
            "statement, the result is 'could not verify', not 'matches'. Green must mean checked against the "
            "bank's own printed figure.",
            "Remember which reader was used for each file. When a CA uses the optional cloud reader for one "
            "stubborn scan, a later re-run must not silently throw that result away and go back to the worse "
            "local result.",
            "Make re-running safe. Today the app deletes all existing rows before re-reading a month; if it "
            "fails halfway, the month ends up emptier than before. Only replace the old rows once the new ones "
            "are ready.",
            "Show the warnings where the CA is looking. Some flags already reach the Excel but never reach the "
            "screen.",
        ],
        [
            "A statement with one deliberately corrupted amount produces a flagged row - not a missing one.",
            "The screen reports how many lines could not be turned into rows, and that number is correct.",
            "A statement with no printed closing balance shows 'could not verify' instead of a green match.",
            "A cloud-read file keeps its rows after the whole month is re-run.",
            "A re-run that fails halfway leaves the month exactly as it was before.",
            "All 126 existing automated tests still pass, plus new ones covering each case above.",
        ],
        "About one week. It touches the bank reader and the balance checker, both of which already have decent "
        "test coverage to work against.",
        "No new banks, no new file types, no reconciliation. This stage only makes the existing output honest.",
        "It is quick, it is contained, and it fixes the one stage that already works on real files - so it is "
        "the safest possible place to start. It also sets the standard the next four stages have to meet: when "
        "the app is unsure, it says so out loud rather than showing a tick it has not earned.",
        PALE_RED,
        "Why this is first:",
    )
    doc.add_page_break()


def build_stage_8(doc):
    stage_section(
        doc,
        "Stage 8",
        "Master reconciliation grid",
        "The main event",
        "F5D89A",
        "Put the client's books, their GST portal data and their bank side by side in one sheet, and show what "
        "does not agree.",
        [
            "Everything up to here converts files. Useful, but a CA could get most of it from a good clerk. "
            "This stage is the one that is actually worth money.",
            "Each month a CA has to check that every purchase invoice in the client's books also appears in the "
            "GSTR-2B the government generated - because that is what decides how much tax credit the client can "
            "claim. Doing it by hand across a few hundred invoices is slow, dull, and easy to get wrong.",
            "This stage does that comparison automatically and produces one sheet that says: these matched, "
            "these are in the books but not on the portal, these are on the portal but not in the books, and "
            "these appear in both but the amounts differ.",
        ],
        [
            "A matching rule. Two entries are the same invoice when the supplier's GSTIN, the invoice number, "
            "the date and the amount all line up, allowing a one-rupee rounding difference.",
            "Sensible near-misses. Invoice numbers are typed inconsistently in the real world, so 'INV/001' and "
            "'INV-1' should be offered as a likely match for the CA to confirm, not declared unmatched.",
            "The output file itself - Master_Reconciliation_Grid.xlsx - with one row per invoice and a clear "
            "status: matched, only in books, only on the portal, or amounts differ.",
            "An on-screen summary, so the CA sees the counts before opening Excel.",
            "A light bank cross-check. Where a payment can be tied to an invoice, show it - clearly marked as a "
            "hint, never as fact.",
            "Every row keeps its trail back to the original document, so clicking a disputed number still shows "
            "the picture of where it came from.",
        ],
        [
            "One folder containing bills or a Tally export plus a GSTR-2B file produces the master grid.",
            "An invoice known to be in both places is reported as matched.",
            "An invoice known to be only on the portal is reported as missing from the books.",
            "An invoice with a deliberate amount difference is reported as an amount clash.",
            "The totals on the grid agree with the totals on the individual registers.",
            "A CA can open the grid in Excel and carry on working in it without cleaning it up first.",
        ],
        "About two to three weeks. This is the largest remaining piece of thinking, mostly because real-world "
        "invoice numbers are messy and the matching rules need testing against genuinely awkward data.",
        "Not the full tax-credit product, not the annual GSTR-9 return, and no automatic filing. This stage "
        "shows the differences. The CA decides what to do about them.",
        "Do not tune the matching rules on invented test files alone. Get one real month from one friendly CA - "
        "with permission - and measure how many invoices match correctly. That number tells you whether this "
        "stage is finished far better than any checklist can.",
        PALE_GOLD,
        "Practical advice:",
    )
    doc.add_page_break()


def build_stage_9(doc):
    stage_section(
        doc,
        "Stage 9",
        "License and paid plans",
        "Making it a business",
        "F5D89A",
        "Let customers pay for the app, without the app ever becoming a place where client documents are stored "
        "online.",
        [
            "The whole promise of this product is that a client's financial documents never leave the CA's "
            "computer. Charging money usually means accounts, servers and uploads - which would break that "
            "promise.",
            "So the design has to keep the two apart: checking whether someone has paid can use the internet. "
            "Reading a client's bank statement never does.",
        ],
        [
            "Two plans to begin with - a starter plan with a monthly file limit, and a full plan without one.",
            "A licence check that talks to a payment provider and nothing else. It sends a licence key. It never "
            "sends a document, a row, or a client name.",
            "A visible, honest file counter, so a CA on the starter plan always knows where they stand rather "
            "than being surprised at the limit.",
            "A clear message at the limit that explains what happened and how to upgrade.",
            "Graceful behaviour with no internet: everything already imported keeps working, and only a brand "
            "new activation is allowed to fail - honestly, with a plain message.",
        ],
        [
            "A test starter licence blocks the file after the limit; a test full licence does not.",
            "With the network off, previously imported months still open and still export.",
            "Inspecting the network traffic during a full month's work shows no document content leaving the PC.",
            "The file counter on screen matches what the licence actually enforces.",
        ],
        "About one to two weeks, most of it careful work rather than difficult work.",
        "No multi-computer syncing, no team accounts, no usage tracking.",
        "Be able to prove the privacy claim, not just state it. A short written note showing exactly what the "
        "licence check sends is worth more to a cautious CA than any marketing page.",
        PALE_BLUE,
        "Worth doing:",
    )
    doc.add_page_break()


def build_stage_10(doc):
    stage_section(
        doc,
        "Stage 10",
        "Firm-ready",
        "Fit for daily use",
        "F5D89A",
        "Make it something a real firm can install once and keep using all year, including cleaning up after "
        "itself.",
        [
            "A tool that works brilliantly in week one and becomes unusable by month three is not a product. "
            "This stage is about everything that only hurts after a while.",
            "The biggest of these is housekeeping. Every file a CA imports is copied into a hidden folder on "
            "their computer and never removed. There is no way to delete a client, a month, or a single file - "
            "the only option is to erase everything and start again. For someone holding other people's "
            "financial records, that is a genuine professional problem, not just untidiness.",
        ],
        [
            "Delete, at every level - one file, one month, one client - removing both the record and the copied "
            "documents from the disk, with a clear warning first.",
            "A one-file installer a non-technical person can run, with no separate software to install.",
            "A visible version number inside the app, so a support conversation can start with a straight answer.",
            "Proper logging. Today each crash overwrites the last one, so there is never a history to look at. "
            "Keep a rolling log the CA can send when something goes wrong.",
            "Remember PDF passwords safely for as long as the app is open, so a protected statement is not asked "
            "about over and over.",
            "Warn when the storage folder sits inside OneDrive or Dropbox, where syncing can corrupt the "
            "database.",
            "Handle a firm with two computers sharing one folder - either safely, or by refusing clearly.",
            "Support older .xls files, which Tally and several bank portals still produce by default.",
            "More bank formats. Five are supported today; the common public sector and newer private banks "
            "should be added from real samples.",
        ],
        [
            "Someone who is not the developer installs from the single file and completes a full bank import.",
            "Deleting a client leaves nothing belonging to that client anywhere on the disk.",
            "The Windows security warning on first install is either removed by signing the app, or has a "
            "written workaround the first customers can follow.",
            "A second computer opening a shared folder either works, or refuses clearly instead of corrupting "
            "data.",
            "A statement from each newly added bank produces a correct register.",
        ],
        "About two to three weeks, and it can overlap with Stage 9 since the two barely touch each other.",
        "No mobile app, no white-labelling, no annual return automation.",
        None,
    )
    doc.add_page_break()


def build_groundwork(doc):
    doc.add_heading("Part 5 - Groundwork that runs alongside", level=1)
    add_body(
        doc,
        "These are not stages. They are things that should be true throughout, and each one quietly makes every "
        "stage after it faster and safer. None of them takes long on its own.",
    )
    add_table(
        doc,
        ["Item", "The problem today", "What to do"],
        [
            (
                "Running the tests",
                "126 tests exist and all pass, but there is no single command that runs them. They have to be "
                "listed out by hand.",
                "Add the one missing file that makes them discoverable, and note the command in the README.",
            ),
            (
                "Automatic checks",
                "Nothing runs the tests or builds the installer automatically. A release once shipped with an "
                "error that stopped the app opening at all.",
                "Set up an automatic check on every change: run the tests, build the installer, confirm the app "
                "starts.",
            ),
            (
                "Crash records",
                "Each crash overwrites the previous one, so there is never any history to look back at.",
                "Keep a rolling log with dates. Add a button that opens it, so a CA can send it.",
            ),
            (
                "Awkward internal code",
                "About a hundred lines in the middle of the app guess at how to call another part of the same "
                "app, trying several combinations until one works.",
                "Fix the connection properly and delete the guesswork. It hides real errors.",
            ),
            (
                "The optional cloud reader",
                "It depends on a separate tool that is not included in the installer, so it only works on the "
                "developer's own machine.",
                "Either include it properly and explain it, or hide the button in the shipped app.",
            ),
        ],
        [1900, 3730, 3730],
    )
    doc.add_page_break()


def build_finish(doc):
    doc.add_heading("What 'finished' actually means", level=1)
    add_body(
        doc,
        "It is worth writing down, because otherwise the work never ends. Version 1.0 is finished when all of "
        "the following are true at the same time:",
    )
    for item in [
        "Every stage from 4 to 7 has a measured accuracy number against real client documents, written down, "
        "and each one is good enough that a CA would not redo the work by hand.",
        "A chartered accountant who has never met the developer can install it from one file and succeed on "
        "their first try, without being talked through it.",
        "They can process a real client month - bank, bills, GST downloads, books - and get a reconciliation "
        "grid they trust enough to work from.",
        "The app never reports a clean result on incomplete data. When it is unsure, it says so.",
        "They can delete a client completely when the engagement ends.",
        "They can pay for it, and can be shown proof that their client documents never left their computer.",
        "When something goes wrong, they can send one file that explains what happened.",
    ]:
        add_check(doc, item)

    add_mini_heading(doc, "The order, one last time")
    add_table(
        doc,
        ["Order", "Stage", "Rough size", "Why here"],
        [
            ("1st", "3.5  Trust repair", "~1 week", "Fixes the one stage proven on real files"),
            ("2nd", "Groundwork (tests, checks)", "~1 week", "A safety net before touching the readers"),
            ("3rd", "6  GST portal files", "~1 week", "The most certain of the four; Stage 8 depends on it"),
            ("4th", "7  Books exports", "~1-2 weeks", "The other half of what Stage 8 compares"),
            ("5th", "4  Trust (scans, crops)", "~1-2 weeks", "Makes everything else verifiable"),
            ("6th", "5  Invoices", "~2-3 weeks", "Hardest and least proven; needs real bills"),
            ("7th", "8  Master grid", "~2-3 weeks", "The feature actually worth paying for"),
            ("8th", "10  Firm-ready", "~2-3 weeks", "Can start alongside; needed before real customers"),
            ("9th", "9  License", "~1-2 weeks", "Only useful once there is something worth buying"),
        ],
        [900, 3000, 1700, 3760],
    )
    add_body(
        doc,
        "That is roughly four to six months of steady solo work. Stages 6 and 7 come before 4 and 5 on "
        "purpose: they are the two Stage 8 depends on, and they are the quickest to make reliable, so the "
        "biggest feature stops being blocked as early as possible.",
        size=10.7,
    )

    add_callout(
        doc,
        "One closing thought:",
        "Stage 9 is listed last on purpose. It is tempting to build payment early because it feels like progress "
        "towards a business. But nobody pays for a converter, and a paywall around an unfinished product is the "
        "fastest way to lose the first customers - who are also the hardest to win back.",
        fill=PALE_BLUE,
    )


def build():
    doc = Document()
    doc.core_properties.title = "CA Unpacker Product Roadmap"
    doc.core_properties.subject = "Plain-English roadmap from working prototype to finished product"
    doc.core_properties.author = "CA Unpacker"
    doc.core_properties.keywords = "CA Unpacker, roadmap, stages, GST, Windows"

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    checklist = styles.add_style("Checklist", 1)
    set_style_font(checklist, "Calibri", 10.7, BLACK)
    checklist.paragraph_format.space_before = Pt(0)
    checklist.paragraph_format.space_after = Pt(4)
    checklist.paragraph_format.line_spacing = 1.25

    hp = section.header.paragraphs[0]
    hp.text = "CA UNPACKER  |  PRODUCT ROADMAP"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MID_GRAY, bold=True)
    add_page_field(section.footer.paragraphs[0])

    build_cover(doc)
    build_how_to_read(doc)
    build_status(doc)
    build_done_part(doc)

    doc.add_heading("Part 2 - Repair first", level=1)
    add_body(
        doc,
        "One short stage, before anything else. It fixes the only stage proven on real documents, and sets the "
        "standard of honesty that the rest of the work has to meet.",
    )
    build_stage_35(doc)

    build_needs_work_part(doc)

    doc.add_heading("Part 4 - The three stages never started", level=1)
    add_body(
        doc,
        "Only open these once Stages 4 to 7 are genuinely reliable. Stage 8 in particular compares those "
        "stages' output against each other - it can only ever be as trustworthy as they are.",
    )
    build_stage_8(doc)
    build_stage_10(doc)
    build_stage_9(doc)

    build_groundwork(doc)
    build_finish(doc)

    doc.save(OUTPUT)
    print("Wrote", OUTPUT)


if __name__ == "__main__":
    build()
