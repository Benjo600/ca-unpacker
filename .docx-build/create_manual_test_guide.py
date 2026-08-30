from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Admin\OneDrive\Desktop\CA idea")
OUTPUT = ROOT / "CA-Unpacker-Prototype-Manual-Test-Guide.docx"
WORKTREE = ROOT / ".worktrees" / "full-product-integration"
TEST_DUMP = WORKTREE / "test-dump"
START_BAT = WORKTREE / "start.bat"

NAVY = "163A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "EAF4EA"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDECEC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "5C6670"
GRID = "CDD5DF"
WHITE = "FFFFFF"
BLACK = "1F2328"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        if edge in kwargs:
            tag = "w:" + edge
            element = tc_mar.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_mar.append(element)
            element.set(qn("w:w"), str(kwargs[edge]))
            element.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=BLACK, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_callout(doc, label, text, fill=PALE_GOLD, color=BLACK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=fill, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + ": ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(value)
    set_run_font(r, size=10.5)


def add_check(doc, text, indent=0.0):
    p = doc.add_paragraph(style="Checklist")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("[ ]  " + text)
    set_run_font(r, size=10.7)
    return p


def add_test(doc, number, title, purpose, steps, expected, pass_rule, notes_lines=2):
    h = doc.add_paragraph(style="Heading 2")
    h.paragraph_format.keep_with_next = True
    r = h.add_run(f"Test {number} - {title}")
    set_run_font(r, size=13, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Purpose: ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(purpose)
    set_run_font(r, size=10.5)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [720, 8640])
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    headers = ("Step", "Action")
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for idx, step in enumerate(steps, start=1):
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(idx))
        set_run_font(r, size=10, color=NAVY, bold=True)
        p = cells[1].paragraphs[0]
        r = p.add_run(step)
        set_run_font(r, size=10)
    set_table_geometry(table, [720, 8640])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Expected result: ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(expected)
    set_run_font(r, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Pass when: ")
    set_run_font(r, size=10.5, color="246B3D", bold=True)
    r = p.add_run(pass_rule)
    set_run_font(r, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Result:  [ ] PASS    [ ] FAIL    [ ] NOT RUN")
    set_run_font(r, size=10.5, bold=True)
    for _ in range(notes_lines):
        p = doc.add_paragraph("Notes: " + "_" * 88)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_run_font(run, size=9, color=MID_GRAY)


def build():
    doc = Document()
    doc.core_properties.title = "CA Unpacker Prototype Manual Test Guide"
    doc.core_properties.subject = "Beginner-friendly manual tests for the Release 0 desktop prototype"
    doc.core_properties.author = "CA Unpacker"
    doc.core_properties.keywords = "CA Unpacker, prototype, manual test, Windows"

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    checklist = styles.add_style("Checklist", 1)
    set_style_font(checklist, "Calibri", 10.7, BLACK)
    checklist.paragraph_format.space_before = Pt(0)
    checklist.paragraph_format.space_after = Pt(4)
    checklist.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "CA UNPACKER  |  PROTOTYPE TEST GUIDE"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MID_GRAY, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_field(fp)

    # Editorial cover header pattern, calibrated for a practical operator guide.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MANUAL TEST GUIDE")
    set_run_font(r, size=11, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CA Unpacker Prototype")
    set_run_font(r, size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A beginner-friendly checklist for testing the Windows desktop app")
    set_run_font(r, size=14, color=DARK_BLUE)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2160, 7200])
    set_table_borders(meta, color="D8E0E8")
    metadata = [
        ("Prepared", "18 August 2026"),
        ("Prototype", "Release 0 - Trust Foundation"),
        ("Estimated time", "15 minutes for the smoke test; 45-60 minutes for the full checklist"),
    ]
    for row, (label, value) in zip(meta.rows, metadata):
        set_cell_shading(row.cells[0], PALE_BLUE)
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=10, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, size=10)
    set_table_geometry(meta, [2160, 7200])
    doc.add_paragraph()
    add_callout(
        doc,
        "Safety first:",
        "Use only the included synthetic files for your first test. Do not use real client documents until you are comfortable with the prototype.",
        fill=PALE_GOLD,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    r = p.add_run("No coding or terminal knowledge is required.")
    set_run_font(r, size=11, color=MID_GRAY, italic=True)
    doc.add_page_break()

    doc.add_heading("Start here", level=1)
    add_callout(
        doc,
        "Important:",
        "This is a source prototype, not the final installer. Run it from the isolated prototype folder described below. Keep the black command window open while testing.",
        fill=PALE_BLUE,
        color=NAVY,
    )
    add_kv(doc, "App launcher", str(START_BAT))
    add_kv(doc, "Synthetic test folder", str(TEST_DUMP))
    add_kv(doc, "Suggested output folder", str(ROOT / "CA Unpacker Test Output"))
    add_kv(doc, "Local app library", r"%LOCALAPPDATA%\CAUnpacker")

    doc.add_heading("Before you begin", level=2)
    add_check(doc, "Close Excel files from any previous CA Unpacker test.")
    add_check(doc, "Create an empty folder named 'CA Unpacker Test Output' in the visible CA idea folder.")
    add_check(doc, "Use a made-up firm and client name. Suggested: Test & Co. and Demo Trading Pvt Ltd.")
    add_check(doc, "Do not click 'Delete library' if this computer already contains data you want to keep.")
    add_check(doc, "Have a place ready for screenshots if anything looks wrong.")

    doc.add_heading("15-minute smoke test", level=1)
    smoke = [
        ("Launch", "Double-click start.bat. The CA Unpacker window should open without closing itself."),
        ("Set up", "On first launch, enter a test firm name, choose the test output folder, then click 'Open the desk'."),
        ("Create work", "Add client 'Demo Trading Pvt Ltd', open it, add period 'Jul 2026', then open the dump tray."),
        ("Import", f"Click 'Add folder' and choose the test-dump folder. It currently contains 13 synthetic files."),
        ("Wait", "Wait until processing stops. Do not close the app while the status is changing."),
        ("Account", "Confirm all 13 filenames appear. Every file must show Processed, Needs review, Failed, or Unclassified."),
        ("Warnings", "Because the mixed folder contains unrelated/unknown files, expect 'Completed with warnings', not a clean success."),
        ("Outputs", "Use 'Open pack folder' and confirm one or more generated Excel files can be opened."),
        ("Persistence", "Close and reopen the app. Confirm the test client, period, files, and outputs are still shown."),
    ]
    smoke_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(smoke_table, [720, 2040, 6600])
    set_table_borders(smoke_table)
    set_repeat_table_header(smoke_table.rows[0])
    for idx, text in enumerate(("Done", "Check", "What should happen")):
        cell = smoke_table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for label, description in smoke:
        cells = smoke_table.add_row().cells
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[ ]")
        set_run_font(r, size=10.5, bold=True)
        r = cells[1].paragraphs[0].add_run(label)
        set_run_font(r, size=10, color=DARK_BLUE, bold=True)
        r = cells[2].paragraphs[0].add_run(description)
        set_run_font(r, size=9.8)
    set_table_geometry(smoke_table, [720, 2040, 6600])
    add_callout(
        doc,
        "Smoke-test pass:",
        "The app launches, imports all 13 files without silently dropping any, shows warning/review outcomes, creates usable Excel output, and remembers the work after restart.",
        fill=PALE_GREEN,
        color="246B3D",
    )

    doc.add_page_break()
    doc.add_heading("Detailed manual tests", level=1)
    p = doc.add_paragraph("Run these after the smoke test. Tick one result for each test and write down anything unexpected.")
    p.paragraph_format.space_after = Pt(10)

    add_test(
        doc,
        1,
        "Launch and first-run setup",
        "Confirm a non-technical user can open and configure the prototype.",
        [
            f"Open {WORKTREE} in File Explorer.",
            "Double-click start.bat and leave the black command window open.",
            "If the 'First launch' screen appears, enter 'Test & Co.' as the firm name.",
            f"Click 'Choose folder' and select {ROOT / 'CA Unpacker Test Output'}.",
            "Click 'Open the desk'. If setup was already completed earlier, simply confirm the Clients screen appears.",
        ],
        "The app window opens and shows the Clients screen. It does not crash, freeze, or upload anything.",
        "The Clients screen is usable and the chosen output folder appears in the left rail.",
    )

    add_test(
        doc,
        2,
        "Create a client and period",
        "Confirm the basic client workspace can be created without special knowledge.",
        [
            "Enter 'Demo Trading Pvt Ltd' in Client name. Leave GSTIN blank or use a clearly fake test value.",
            "Click 'Add client', then click the new client row.",
            "Enter 'Jul 2026' in Period and click 'Add period'.",
            "Click the new period row labelled 'Open dump tray'.",
        ],
        "The dump tray opens and its heading identifies Demo Trading Pvt Ltd and Jul 2026.",
        "The client and period appear once, with no error message or duplicate row.",
    )

    add_test(
        doc,
        3,
        "Import the mixed synthetic folder",
        "Confirm mixed PDFs, JSON, XML, ZIP, CSV, images, and an unrelated document are all accounted for.",
        [
            "In the dump tray, click 'Add folder'.",
            f"Choose {TEST_DUMP} and confirm the folder selection.",
            "Wait until the status stops changing and the Add buttons become available again.",
            "Count the filenames shown in the file ledger. Compare them with the 13-file checklist in Appendix A.",
        ],
        "Exactly 13 files are visible. Each one has a type and one terminal outcome: Processed, Needs review, Failed, or Unclassified.",
        "No selected file disappears, processing finishes, and the visible outcome count reconciles to 13.",
    )

    doc.add_page_break()
    add_test(
        doc,
        4,
        "Review warnings and unknown files",
        "Confirm the app does not claim clean success when files still need attention.",
        [
            "Read the status at the top-right of the dump tray.",
            "Open the 'Review files' section.",
            "Find meeting_notes.docx and random_scan.jpg, or any file marked Needs review, Failed, or Unclassified.",
            "Confirm each review item shows a reason instead of only a colored badge.",
        ],
        "The job says 'Completed with warnings' when unresolved files exist. Review items show the source filename, outcome, and a reason.",
        "The app never shows clean completion while unresolved or failed files are present.",
    )

    add_test(
        doc,
        5,
        "Open and spot-check Excel outputs",
        "Confirm generated workbooks exist, open in Excel, and contain recognizable data from the synthetic sources.",
        [
            "Click 'Open pack folder'. If 'Open in Excel' is available, test that button too.",
            "Look for the workbook names listed in Appendix B. The exact set depends on which sample files processed successfully.",
            "Open each available workbook and check that its first row contains headings rather than raw PDF text.",
            "For Bank_Statement_Cleaned.xlsx, confirm dates/amounts are present and no obviously unrelated invoice or GST rows appear.",
            "Close the workbooks before continuing so files are not locked.",
        ],
        "Available .xlsx files open normally, contain structured rows, and are stored under the chosen test output folder.",
        "At least one expected workbook opens without a repair warning, and its contents match the document type named in the workbook.",
    )

    add_test(
        doc,
        6,
        "Manual type override and reprocessing",
        "Confirm a reviewer can correct a classification and the app reprocesses the period transparently.",
        [
            "In Review files, choose an unresolved test item such as random_scan.jpg.",
            "Change its type using the dropdown. Use the most reasonable option shown, such as Invoice.",
            "Wait for automatic reprocessing to finish.",
            "Confirm the file remains visible and receives a new terminal outcome and reason.",
            "If the override is clearly wrong, change it back to Unknown after recording the result.",
        ],
        "The app becomes busy, reparses the current period, and returns to a terminal job status without losing the file or the other outputs.",
        "The override is visible, reprocessing finishes, and every file still has an outcome.",
    )

    add_test(
        doc,
        7,
        "Reject legacy .xls safely",
        "Confirm unsupported Excel files are rejected with useful guidance rather than accepted silently.",
        [
            f"Create a new empty folder named 'xls-rejection-test' under {ROOT}.",
            "Open Notepad, type the word test, choose Save As, select 'All files', and save it inside that folder as legacy-test.xls.",
            "In CA Unpacker, click 'Add folder' and select xls-rejection-test.",
            "Read the error shown in the dump tray.",
        ],
        "The import is rejected and the message says legacy .xls is unsupported and should be exported as .xlsx or .csv.",
        "No false Excel output is produced and the guidance is understandable.",
    )

    add_test(
        doc,
        8,
        "Restart and persistence",
        "Confirm local work survives a normal close and reopen.",
        [
            "Write down the visible client, period, file count, and one output filename.",
            "Close the CA Unpacker window, then close the black command window if it remains open.",
            "Double-click start.bat again.",
            "Open Demo Trading Pvt Ltd and Jul 2026.",
            "Compare the restored files and outputs with your notes.",
        ],
        "The client, period, imported files, outcomes, and generated pack remain available after restart.",
        "Nothing has to be imported again and the stored file count matches the pre-restart count.",
    )

    doc.add_heading("What not to test manually yet", level=1)
    add_callout(
        doc,
        "Prototype boundary:",
        "The following gates are important, but they are better handled by automated tests or a later installed-app build. Skipping them does not invalidate your manual prototype feedback.",
        fill=LIGHT_GRAY,
    )
    items = [
        "The 400/401-file folder boundary. Automated tests cover this and creating 401 files by hand adds little value.",
        "Password-protected bank PDFs unless you already have a safe synthetic sample.",
        "A clean Windows 11 installer test. This source prototype is not the final packaged installer.",
        "Tesseract packaging on a clean computer. OCR may vary until the release build bundles the OCR assets.",
        "Real client documents. Keep them outside Git and wait until the prototype workflow feels trustworthy.",
    ]
    for item in items:
        add_check(doc, item)

    doc.add_heading("How to report a problem", level=1)
    p = doc.add_paragraph("For every failure, capture these five things. This is enough information to reproduce most prototype problems.")
    report_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(report_table, [2160, 7200])
    set_table_borders(report_table)
    set_repeat_table_header(report_table.rows[0])
    for idx, header in enumerate(("Record", "What to write")):
        set_cell_shading(report_table.rows[0].cells[idx], PALE_BLUE)
        r = report_table.rows[0].cells[idx].paragraphs[0].add_run(header)
        set_run_font(r, size=10, color=NAVY, bold=True)
    rows = [
        ("Test number", "For example: Test 4 - Review warnings and unknown files"),
        ("Exact filename", "Copy the full name shown in the app"),
        ("Expected", "What this guide said should happen"),
        ("Actual", "What happened instead, including the exact error wording"),
        ("Evidence", "A screenshot plus the last visible lines in the black command window"),
    ]
    for label, value in rows:
        cells = report_table.add_row().cells
        r = cells[0].paragraphs[0].add_run(label)
        set_run_font(r, size=10, color=DARK_BLUE, bold=True)
        r = cells[1].paragraphs[0].add_run(value)
        set_run_font(r, size=10)
    set_table_geometry(report_table, [2160, 7200])
    add_callout(
        doc,
        "Stop immediately if:",
        "a selected file vanishes, a failed file produces clean completion, the app overwrites an original, or an unexpected network/upload message appears.",
        fill=PALE_RED,
        color="9B1C1C",
    )

    doc.add_heading("Test session summary", level=1)
    summary = doc.add_table(rows=1, cols=4)
    set_table_geometry(summary, [720, 2400, 1440, 4800])
    set_table_borders(summary)
    set_repeat_table_header(summary.rows[0])
    for idx, header in enumerate(("Test", "Area", "Result", "Notes / screenshot name")):
        set_cell_shading(summary.rows[0].cells[idx], PALE_BLUE)
        p = summary.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    areas = [
        "Launch/setup", "Client/period", "Folder import", "Warnings/review",
        "Excel outputs", "Override/reparse", ".xls rejection", "Restart/persistence",
    ]
    for idx, area in enumerate(areas, start=1):
        cells = summary.add_row().cells
        values = (str(idx), area, "P / F / NR", "")
        for col, value in enumerate(values):
            p = cells[col].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=MID_GRAY if col == 3 else BLACK)
    set_table_geometry(summary, [720, 2400, 1440, 4800])
    add_kv(doc, "Tester", "________________________________________")
    add_kv(doc, "Date and time", "________________________________________")
    add_kv(doc, "Overall verdict", "[ ] Ready for another feedback round   [ ] Fix issues first")

    doc.add_page_break()
    doc.add_heading("Appendix A - Synthetic test files", level=1)
    p = doc.add_paragraph(f"Folder: {TEST_DUMP}")
    p.paragraph_format.space_after = Pt(8)
    files = sorted([path.name for path in TEST_DUMP.iterdir() if path.is_file()], key=str.lower)
    file_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(file_table, [720, 5040, 3600])
    set_table_borders(file_table)
    set_repeat_table_header(file_table.rows[0])
    for idx, header in enumerate(("Seen", "Filename", "Likely family")):
        set_cell_shading(file_table.rows[0].cells[idx], PALE_BLUE)
        r = file_table.rows[0].cells[idx].paragraphs[0].add_run(header)
        set_run_font(r, size=10, color=NAVY, bold=True)
    family_map = {
        ".json": "GST return data", ".pdf": "Bank statement or invoice", ".png": "Scanned/image invoice",
        ".jpg": "Image requiring classification", ".docx": "Unrelated/unknown document", ".zip": "Tally backup",
        ".xml": "Tally export", ".csv": "Books/Zoho export",
    }
    for filename in files:
        cells = file_table.add_row().cells
        suffix = Path(filename).suffix.lower()
        values = ("[ ]", filename, family_map.get(suffix, "Unknown"))
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=9.7)
    set_table_geometry(file_table, [720, 5040, 3600])

    doc.add_page_break()
    doc.add_heading("Appendix B - Expected workbook names", level=1)
    p = doc.add_paragraph("You may see some or all of these, depending on which synthetic inputs processed successfully:")
    workbook_names = [
        "Bank_Statement_Cleaned.xlsx",
        "Purchase_Register_Extracted.xlsx",
        "Sales_Register_Extracted.xlsx",
        "GSTR_1_Formatted.xlsx",
        "GSTR_2B_Formatted.xlsx",
        "GSTR_3B_Formatted.xlsx",
        "Books_Register_Extracted.xlsx",
    ]
    for name in workbook_names:
        add_check(doc, name)

    # Apply keep-together behavior to compact note lines and ensure no table row is fixed-height.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
